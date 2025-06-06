import os
import pypdf
import zipfile
import tempfile
import shutil
from typing import Type
from pydantic import BaseModel, Field
from crewai.tools import BaseTool
import logging
import PyPDF2
import re

logger = logging.getLogger(__name__)


class SimpleFileReadToolInput(BaseModel):
    """Input schema for SimpleFileReadTool."""

    file_path: str = Field(..., description="Path to the file to be read.")
    max_chars: int = Field(20000, description="Maximum number of characters to return.")


class SimpleFileReadTool(BaseTool):
    name: str = "Simple File Reader"
    description: str = (
        "Reads the content of a file and returns it as text. "
        "Supports TXT, PDF, DOCX, MD files and ZIP archives."
    )
    args_schema: Type[BaseModel] = SimpleFileReadToolInput

    def _run(self, file_path: str, max_chars: int = 20000) -> str:
        """Read file content and return it as text."""
        try:
            logger.info(f"Tentando ler arquivo: {file_path}")
            logger.info(f"Caminho absoluto atual: {os.path.abspath('.')}")
            
            # Se o caminho for relativo, tenta encontrar o arquivo no diretório atual
            if not os.path.isabs(file_path):
                # Tenta encontrar o arquivo no diretório atual
                current_dir = os.getcwd()
                possible_paths = [
                    os.path.join(current_dir, file_path),
                    os.path.join(current_dir, "samples", file_path),
                    os.path.join(current_dir, "samples", "edital-004", file_path),
                    os.path.join(current_dir, "edital-004", file_path),
                    os.path.join("/mnt/c/Users/SDS/Documents/edital-summarizer/samples/edital-004", file_path)
                ]
                
                logger.info(f"Procurando arquivo em: {possible_paths}")
                
                for path in possible_paths:
                    if os.path.exists(path):
                        file_path = path
                        logger.info(f"Arquivo encontrado em: {file_path}")
                        break
                else:
                    logger.error(f"Arquivo não encontrado em nenhum dos caminhos possíveis: {possible_paths}")
                    return f"Error: File not found at {file_path}"

            logger.info(f"Tentando ler arquivo: {file_path}")
            
            if not os.path.exists(file_path):
                logger.error(f"Arquivo não encontrado: {file_path}")
                return f"Error: File not found at {file_path}"

            _, file_extension = os.path.splitext(file_path)
            file_extension = file_extension.lower()

            try:
                if file_extension == ".pdf":
                    text = self._extract_text_from_pdf(file_path)
                elif file_extension in [".docx", ".doc"]:
                    text = self._extract_text_from_docx(file_path)
                elif file_extension in [".md", ".markdown"]:
                    text = self._extract_text_from_markdown(file_path)
                elif file_extension == ".zip":
                    text = self._extract_text_from_zip(file_path, max_chars)
                else:  # Assume it's a text file
                    with open(file_path, "r", encoding="utf-8", errors="replace") as file:
                        text = file.read()

                # Limit text size
                if len(text) > max_chars:
                    text = (
                        text[:max_chars] + f"\n\n[Texto truncado em {max_chars} caracteres]"
                    )

                logger.info(f"Arquivo lido com sucesso: {file_path}")
                logger.debug(f"Primeiros 100 caracteres do arquivo: {text[:100]}")
                return text
            except Exception as e:
                logger.error(f"Erro ao ler arquivo {file_path}: {str(e)}")
                return f"Error reading file: {str(e)}"
        except Exception as e:
            logger.error(f"Erro ao processar arquivo {file_path}: {str(e)}")
            return f"Error processing file: {str(e)}"

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            logger.info(f"Extraindo texto do PDF: {file_path}")
            text = ""
            
            with open(file_path, "rb") as file:
                # Tenta ler o PDF com PyPDF2
                try:
                    pdf_reader = PyPDF2.PdfReader(file)
                    num_pages = len(pdf_reader.pages)
                    logger.info(f"PDF tem {num_pages} páginas")
                    
                    # Processa apenas as primeiras 20 páginas para economizar tokens
                    max_pages = min(20, num_pages)
                    logger.info(f"Processando {max_pages} páginas do PDF")
                    
                    for page_num in range(max_pages):
                        try:
                            page = pdf_reader.pages[page_num]
                            page_text = page.extract_text()
                            if page_text:
                                # Limpa o texto da página antes de adicionar
                                page_text = page_text.replace('\x00', '')  # Remove caracteres nulos
                                page_text = re.sub(r'\s+', ' ', page_text)  # Normaliza espaços
                                page_text = page_text.strip()
                                
                                # Adiciona apenas se houver conteúdo significativo
                                if len(page_text) > 50:  # Ignora páginas com pouco conteúdo
                                    text += f"\n\n=== Página {page_num + 1} ===\n\n{page_text}"
                                    logger.debug(f"Texto extraído da página {page_num + 1}: {page_text[:100]}")
                                else:
                                    logger.warning(f"Página {page_num + 1} ignorada por ter pouco conteúdo")
                            else:
                                logger.warning(f"Nenhum texto extraído da página {page_num + 1}")
                        except Exception as e:
                            logger.error(f"Erro ao extrair texto da página {page_num + 1}: {str(e)}")
                            continue
                except Exception as e:
                    logger.error(f"Erro ao ler PDF com PyPDF2: {str(e)}")
                    # Se falhar com PyPDF2, tenta com pypdf
                    try:
                        file.seek(0)  # Volta ao início do arquivo
                        pdf_reader = pypdf.PdfReader(file)
                        num_pages = len(pdf_reader.pages)
                        logger.info(f"PDF tem {num_pages} páginas (usando pypdf)")
                        
                        # Processa apenas as primeiras 20 páginas para economizar tokens
                        max_pages = min(20, num_pages)
                        logger.info(f"Processando {max_pages} páginas do PDF")
                        
                        for page_num in range(max_pages):
                            try:
                                page = pdf_reader.pages[page_num]
                                page_text = page.extract_text()
                                if page_text:
                                    # Limpa o texto da página antes de adicionar
                                    page_text = page_text.replace('\x00', '')  # Remove caracteres nulos
                                    page_text = re.sub(r'\s+', ' ', page_text)  # Normaliza espaços
                                    page_text = page_text.strip()
                                    
                                    # Adiciona apenas se houver conteúdo significativo
                                    if len(page_text) > 50:  # Ignora páginas com pouco conteúdo
                                        text += f"\n\n=== Página {page_num + 1} ===\n\n{page_text}"
                                        logger.debug(f"Texto extraído da página {page_num + 1}: {page_text[:100]}")
                                    else:
                                        logger.warning(f"Página {page_num + 1} ignorada por ter pouco conteúdo")
                                else:
                                    logger.warning(f"Nenhum texto extraído da página {page_num + 1}")
                            except Exception as e:
                                logger.error(f"Erro ao extrair texto da página {page_num + 1}: {str(e)}")
                                continue
                    except Exception as e:
                        logger.error(f"Erro ao ler PDF com pypdf: {str(e)}")
                        raise
            
            if not text.strip():
                logger.error("Nenhum texto extraído do PDF")
                return "Error: No text extracted from PDF"
            
            # Limpa o texto extraído
            text = text.replace('\x00', '')  # Remove caracteres nulos
            text = re.sub(r'\s+', ' ', text)  # Normaliza espaços
            text = text.strip()
            
            # Limita o tamanho do texto para economizar tokens
            max_chars = 20000  # Limite máximo de caracteres
            if len(text) > max_chars:
                text = text[:max_chars] + "\n\n[Texto truncado]"
            
            logger.info(f"Texto extraído com sucesso do PDF. Tamanho: {len(text)} caracteres")
            return text
            
        except Exception as e:
            logger.error(f"Erro ao extrair texto do PDF: {str(e)}")
            return f"Error: Failed to extract text from PDF: {str(e)}"

    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            logger.info(f"Extraindo texto do DOCX: {file_path}")
            text = ""
            
            import docx
            doc = docx.Document(file_path)
            
            # Extrai texto dos parágrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extrai texto das tabelas
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            if not text.strip():
                logger.error("Nenhum texto extraído do DOCX")
                return "Error: No text extracted from DOCX"
            
            logger.info(f"Texto extraído com sucesso do DOCX. Tamanho: {len(text)} caracteres")
            return text
            
        except Exception as e:
            logger.error(f"Erro ao extrair texto do DOCX: {str(e)}")
            return f"Error: Failed to extract text from DOCX: {str(e)}"

    def _extract_text_from_markdown(self, file_path: str) -> str:
        """Extract text from Markdown file."""
        try:
            logger.info(f"Extraindo texto do Markdown: {file_path}")
            
            with open(file_path, "r", encoding="utf-8", errors="replace") as file:
                text = file.read()
            
            if not text.strip():
                logger.error("Nenhum texto extraído do Markdown")
                return "Error: No text extracted from Markdown"
            
            logger.info(f"Texto extraído com sucesso do Markdown. Tamanho: {len(text)} caracteres")
            return text
            
        except Exception as e:
            logger.error(f"Erro ao extrair texto do Markdown: {str(e)}")
            return f"Error: Failed to extract text from Markdown: {str(e)}"

    def _extract_text_from_zip(self, file_path: str, max_chars: int) -> str:
        """Extract text from ZIP file."""
        try:
            logger.info(f"Extraindo texto do ZIP: {file_path}")
            text = ""
            temp_dir = None
            
            try:
                # Cria um diretório temporário
                temp_dir = tempfile.mkdtemp()
                logger.info(f"Diretório temporário criado: {temp_dir}")
                
                # Extrai o arquivo ZIP
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)
                logger.info("Arquivo ZIP extraído com sucesso")
                
                # Processa os arquivos extraídos
                for root, _, files in os.walk(temp_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        logger.info(f"Processando arquivo extraído: {file_path}")
                        
                        try:
                            # Tenta extrair texto do arquivo
                            file_tool = SimpleFileReadTool()
                            file_text = file_tool._run(file_path)
                            
                            if not file_text.startswith("Error:"):
                                text += f"\n\n=== {file} ===\n\n{file_text}"
                                logger.debug(f"Texto extraído do arquivo {file}: {file_text[:100]}")
                            else:
                                logger.warning(f"Erro ao extrair texto do arquivo {file}: {file_text}")
                                
                        except Exception as e:
                            logger.error(f"Erro ao processar arquivo {file}: {str(e)}")
                            continue
                
                if not text.strip():
                    logger.error("Nenhum texto extraído do ZIP")
                    return "Error: No text extracted from ZIP"
                
                # Limita o tamanho do texto
                if len(text) > max_chars:
                    text = text[:max_chars] + f"\n\n[Texto truncado em {max_chars} caracteres]"
                
                logger.info(f"Texto extraído com sucesso do ZIP. Tamanho: {len(text)} caracteres")
                return text
                
            finally:
                # Limpa o diretório temporário
                if temp_dir and os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
                    logger.info(f"Diretório temporário removido: {temp_dir}")
            
        except Exception as e:
            logger.error(f"Erro ao extrair texto do ZIP: {str(e)}")
            return f"Error: Failed to extract text from ZIP: {str(e)}"
