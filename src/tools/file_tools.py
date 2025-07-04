import os
import pypdf
import zipfile
import tempfile
import shutil
from typing import Type
from pydantic import BaseModel, Field
from crewai.tools import BaseTool
import PyPDF2
import re
import docx
import json
import csv
from utils import get_logger

logger = get_logger(__name__)

"""
FileReadTool - Ferramenta para leitura de arquivos de editais

Suporte a arquivos compactados:
- ZIPs aninhados (até 3 níveis de profundidade)
- Detecção automática de arquivos compactados
- Processamento recursivo seguro com limite de profundidade
- Limpeza automática de diretórios temporários

Formatos suportados:
- PDF, DOCX, TXT, MD, CSV, JSON
- ZIP (com suporte a aninhamento)
- Outros formatos compactados detectados mas não processados recursivamente
"""

def clean_text(text: str) -> str:
    """Limpa e otimiza o texto removendo espaços desnecessários e caracteres especiais."""
    if not text:
        return text
        
    # Remove caracteres nulos e outros caracteres especiais
    text = text.replace('\x00', '')
    text = text.replace('\r', ' ')
    
    # Remove múltiplos espaços em branco
    text = re.sub(r'\s+', ' ', text)
    
    # Remove espaços antes e depois de pontuação
    text = re.sub(r'\s+([.,;:!?])', r'\1', text)
    
    # Remove quebras de linha desnecessárias
    text = re.sub(r'\n\s*\n', '\n', text)
    
    # Remove espaços no início e fim de cada linha
    text = '\n'.join(line.strip() for line in text.split('\n'))
    
    # Remove linhas vazias no início e fim do texto
    text = text.strip()
    
    return text

class DocumentTooLargeError(Exception):
    """Exceção lançada quando o documento excede o limite de caracteres."""
    def __init__(self, max_chars: int, actual_chars: int):
        self.max_chars = max_chars
        self.actual_chars = actual_chars
        self.error_message = (
            f"Não foi possível processar a análise por completo pois o documento é muito grande "
            f"(tamanho atual: {actual_chars} caracteres, limite: {max_chars} caracteres). "
            f"Por segurança, o edital foi marcado como não relevante."
        )
        super().__init__(self.error_message)

class InsufficientContentError(Exception):
    """Exceção lançada quando não há conteúdo suficiente para análise."""
    def __init__(self, error_message: str):
        self.error_message = error_message
        super().__init__(self.error_message)

class FileReadToolInput(BaseModel):
    """Input schema for FileReadTool."""
    edital_path_dir: str = Field(..., description="Path to the edital directory containing the files to be read.")
    max_chars: int = Field(100000, description="Maximum number of characters to return.")

class FileReadTool(BaseTool):
    """Tool for reading different file formats."""
    
    name: str = "File Reader"
    description: str = (
        "Reads the content of all files in an edital directory and returns it as text. "
        "Supports TXT, PDF, DOCX, MD, CSV, JSON files and ZIP archives (including nested ZIPs up to 3 levels deep)."
    )
    args_schema: Type[BaseModel] = FileReadToolInput

    def _run(self, edital_path_dir: str, max_chars: int = 200000) -> str:
        """Read all files in the edital directory and return their content as text."""
        try:
            print(f"\n{'='*50}")
            print(f"FileReadTool: Iniciando leitura exploratória do diretório: {edital_path_dir}")
            print(f"FileReadTool: Limite de caracteres: {max_chars}")
            
            # Valida caminho
            if not os.path.exists(edital_path_dir):
                print(f"❌ FileReadTool: Diretório não encontrado: {edital_path_dir}")
                raise FileNotFoundError(f"Diretório não encontrado: {edital_path_dir}")

            print(f"✅ FileReadTool: Diretório encontrado")

            # Lista todos os arquivos no diretório
            all_files = []
            metadata_files = []
            content_files = []
            
            for root, _, files in os.walk(edital_path_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    if file.lower() == "metadata.json":
                        metadata_files.append(file_path)
                    else:
                        content_files.append(file_path)
                    all_files.append(file_path)

            print(f"FileReadTool: Encontrados {len(all_files)} arquivos para processar")
            print(f"FileReadTool: - Metadata files: {len(metadata_files)}")
            print(f"FileReadTool: - Content files: {len(content_files)}")

            # Processa cada arquivo
            full_text = ""
            processed_files = 0
            failed_files = []
            successful_content_files = 0
            
            for file_path in all_files:
                try:
                    print(f"\nFileReadTool: Processando arquivo: {file_path}")
                    
                    # Extrai texto baseado no tipo de arquivo
                    file_text = self._extract_text_from_file(file_path)
                    
                    if file_text and not file_text.startswith("Error:"):
                        # Adiciona cabeçalho com informações do arquivo
                        file_name = os.path.basename(file_path)
                        file_header = f"\n\n=== {file_name} ===\n\n"
                        
                        full_text += file_header + file_text
                        processed_files += 1
                        
                        # Conta arquivos de conteúdo processados com sucesso
                        if file_name.lower() != "metadata.json":
                            successful_content_files += 1
                        
                        print(f"✅ FileReadTool: Arquivo processado com sucesso")
                    else:
                        # Só conta como falha se não for metadata.json
                        file_name = os.path.basename(file_path)
                        if file_name.lower() != "metadata.json":
                            failed_files.append(file_name)
                        print(f"❌ FileReadTool: Erro ao processar arquivo: {file_text}")
                        
                except Exception as e:
                    # Só conta como falha se não for metadata.json
                    file_name = os.path.basename(file_path)
                    if file_name.lower() != "metadata.json":
                        failed_files.append(file_name)
                    print(f"❌ FileReadTool: Erro ao processar arquivo {file_path}: {str(e)}")
                    continue

            print(f"\nFileReadTool: Resumo do processamento:")
            print(f"FileReadTool: - Arquivos processados com sucesso: {processed_files}")
            print(f"FileReadTool: - Arquivos de conteúdo processados com sucesso: {successful_content_files}")
            print(f"FileReadTool: - Arquivos de conteúdo com erro: {len(failed_files)}")
            if failed_files:
                print(f"FileReadTool: - Arquivos de conteúdo com erro: {failed_files}")

            # Verifica se há conteúdo suficiente para análise
            if not full_text.strip():
                print("❌ FileReadTool: Nenhum texto extraído dos arquivos")
                return "Error: No text extracted from files"
            
            # Verifica se há apenas metadata.json (sem conteúdo real)
            if len(content_files) == 0:
                error_msg = "Apenas metadata.json encontrado. Não há arquivos de conteúdo para análise."
                print(f"❌ FileReadTool: {error_msg}")
                raise InsufficientContentError(error_msg)
            
            # Verifica se todos os arquivos de conteúdo falharam
            if successful_content_files == 0:
                error_msg = f"Não foi possível extrair conteúdo de nenhum arquivo de conteúdo. Arquivos com erro: {', '.join(failed_files)}"
                print(f"❌ FileReadTool: {error_msg}")
                raise InsufficientContentError(error_msg)

            # Limpa e otimiza o texto
            print("FileReadTool: Limpando e otimizando texto...")
            original_size = len(full_text)
            full_text = clean_text(full_text)
            cleaned_size = len(full_text)
            print(f"FileReadTool: Tamanho original: {original_size} caracteres")
            print(f"FileReadTool: Tamanho após limpeza: {cleaned_size} caracteres")
            print(f"FileReadTool: Redução: {original_size - cleaned_size} caracteres ({(original_size - cleaned_size)/original_size*100:.1f}%)")

            # Verifica limite de caracteres
            if len(full_text) > max_chars:
                print(f"❌ FileReadTool: Documento excede limite de caracteres")
                print(f"FileReadTool: Tamanho atual: {len(full_text)}")
                print(f"FileReadTool: Limite: {max_chars}")
                error_msg = (
                    f"Não foi possível processar a análise por completo pois o documento é muito grande "
                    f"(tamanho atual: {len(full_text)} caracteres, limite: {max_chars} caracteres). "
                    f"Por segurança, o edital foi marcado como não relevante."
                )
                raise DocumentTooLargeError(max_chars, len(full_text))

            print(f"✅ FileReadTool: Processamento concluído com sucesso")
            print(f"{'='*50}\n")
            return full_text
                
        except Exception as e:
            print(f"❌ FileReadTool: Erro ao processar diretório")
            print(f"FileReadTool: {str(e)}")
            raise Exception(f"Erro ao processar diretório: {str(e)}")

    def _extract_text_from_file(self, file_path: str) -> str:
        """Extract text from a file based on its extension."""
        try:
            _, file_extension = os.path.splitext(file_path)
            file_extension = file_extension.lower()

            if file_extension == ".pdf":
                return self._extract_text_from_pdf(file_path)
            elif file_extension in [".docx", ".doc"]:
                return self._extract_text_from_docx(file_path)
            elif file_extension in [".md", ".markdown"]:
                return self._extract_text_from_markdown(file_path)
            elif file_extension == ".zip":
                return self._extract_text_from_zip(file_path, 50000)  # Limite específico para ZIP
            elif file_extension == ".csv":
                return self._extract_text_from_csv(file_path)
            elif file_extension == ".json":
                return self._extract_text_from_json(file_path)
            else:  # Assume it's a text file
                return self._extract_text_from_text_file(file_path)

        except Exception as e:
            return f"Error: Failed to extract text from file: {str(e)}"

    def _extract_text_from_text_file(self, file_path: str) -> str:
        """Extract text from a text file with multiple encoding support."""
        try:
            print(f"FileReadTool: Extraindo texto do arquivo de texto: {file_path}")
            
            # Lista de encodings para tentar (incluindo UTF-8 com BOM)
            encodings = ['utf-8-sig', 'utf-8', 'latin1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        text = file.read()
                    
                    if not text.strip():
                        print(f"FileReadTool: Nenhum texto extraído do arquivo com encoding {encoding}")
                        continue
                    
                    print(f"FileReadTool: Texto extraído com sucesso do arquivo usando encoding {encoding}. Tamanho: {len(text)} caracteres")
                    return text
                    
                except (UnicodeDecodeError, Exception) as e:
                    print(f"FileReadTool: Erro ao ler arquivo com encoding {encoding}: {str(e)}")
                    continue
            
            print("FileReadTool: Não foi possível ler o arquivo com nenhum encoding")
            return "Error: Failed to extract text from file: Could not decode with any encoding"
            
        except Exception as e:
            print(f"FileReadTool: Erro ao extrair texto do arquivo: {str(e)}")
            return f"Error: Failed to extract text from file: {str(e)}"

    def _extract_text_from_csv(self, file_path: str) -> str:
        """Extract text from CSV file."""
        try:
            print(f"FileReadTool: Extraindo texto do CSV: {file_path}")
            
            # Lista de encodings para tentar (incluindo UTF-8 com BOM)
            encodings = ['utf-8-sig', 'utf-8', 'latin1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    text = []
                    
                    with open(file_path, "r", encoding=encoding) as file:
                        csv_reader = csv.reader(file)
                        for row in csv_reader:
                            if any(cell.strip() for cell in row):  # Ignora linhas vazias
                                text.append(" | ".join(cell.strip() for cell in row))
                    
                    if not text:
                        print(f"FileReadTool: Nenhum texto extraído do CSV com encoding {encoding}")
                        continue
                    
                    result = "\n".join(text)
                    print(f"FileReadTool: Texto extraído com sucesso do CSV usando encoding {encoding}. Tamanho: {len(result)} caracteres")
                    return result
                    
                except (UnicodeDecodeError, Exception) as e:
                    print(f"FileReadTool: Erro ao ler CSV com encoding {encoding}: {str(e)}")
                    continue
            
            print("FileReadTool: Não foi possível ler o CSV com nenhum encoding")
            return "Error: Failed to extract text from CSV: Could not decode with any encoding"
            
        except Exception as e:
            print(f"FileReadTool: Erro ao extrair texto do CSV: {str(e)}")
            return f"Error: Failed to extract text from CSV: {str(e)}"

    def _extract_text_from_json(self, file_path: str) -> str:
        """Extract text from JSON file."""
        try:
            print(f"FileReadTool: Extraindo texto do JSON: {file_path}")
            
            # Lista de encodings para tentar (incluindo UTF-8 com BOM)
            encodings = ['utf-8-sig', 'utf-8', 'latin1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        data = json.load(file)
                    
                    # Remove campos específicos se for metadata.json
                    file_name = os.path.basename(file_path).lower()
                    if file_name == "metadata.json":
                        # Remove campos que não devem ser incluídos na análise
                        if "threshold" in data:
                            del data["threshold"]
                            print(f"FileReadTool: Campo 'threshold' removido do metadata.json")
                        if "target" in data:
                            del data["target"]
                            print(f"FileReadTool: Campo 'target' removido do metadata.json")
                    
                    # Converte o JSON em texto formatado
                    text = json.dumps(data, ensure_ascii=False, indent=2)
                    
                    # Para metadata.json, retorna o conteúdo mesmo que esteja vazio
                    if file_name == "metadata.json":
                        print(f"FileReadTool: Metadata.json processado com sucesso usando encoding {encoding}. Tamanho: {len(text)} caracteres")
                        return text
                    
                    # Para outros arquivos JSON, verifica se há conteúdo
                    if not text.strip():
                        print(f"FileReadTool: Nenhum texto extraído do JSON com encoding {encoding}")
                        continue
                    
                    print(f"FileReadTool: Texto extraído com sucesso do JSON usando encoding {encoding}. Tamanho: {len(text)} caracteres")
                    return text
                    
                except (UnicodeDecodeError, json.JSONDecodeError) as e:
                    print(f"FileReadTool: Erro ao ler com encoding {encoding}: {str(e)}")
                    continue
            
            print("FileReadTool: Não foi possível ler o JSON com nenhum encoding")
            return "Error: Failed to extract text from JSON: Could not decode with any encoding"
            
        except Exception as e:
            print(f"FileReadTool: Erro ao extrair texto do JSON: {str(e)}")
            return f"Error: Failed to extract text from JSON: {str(e)}"

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            print(f"FileReadTool: Extraindo texto do PDF: {file_path}")
            text = ""
            
            with open(file_path, "rb") as file:
                # Tenta ler o PDF com PyPDF2
                try:
                    pdf_reader = PyPDF2.PdfReader(file)
                    num_pages = len(pdf_reader.pages)
                    print(f"FileReadTool: PDF tem {num_pages} páginas")
                    
                    # Processa apenas as primeiras 20 páginas
                    max_pages = min(20, num_pages)
                    print(f"FileReadTool: Processando {max_pages} páginas do PDF")
                    
                    for page_num in range(max_pages):
                        try:
                            page = pdf_reader.pages[page_num]
                            page_text = page.extract_text()
                            if page_text:
                                # Limpa o texto da página
                                page_text = clean_text(page_text)
                                
                                # Adiciona apenas se houver conteúdo significativo
                                if len(page_text) > 50:
                                    text += f"\n\n=== Página {page_num + 1} ===\n\n{page_text}"
                                    print(f"FileReadTool: Texto extraído da página {page_num + 1}: {len(page_text)} caracteres")
                                else:
                                    print(f"FileReadTool: Página {page_num + 1} ignorada por ter pouco conteúdo")
                            else:
                                print(f"FileReadTool: Nenhum texto extraído da página {page_num + 1}")
                        except Exception as e:
                            print(f"FileReadTool: Erro ao extrair texto da página {page_num + 1}: {str(e)}")
                            continue
                except Exception as e:
                    print(f"FileReadTool: Erro ao ler PDF com PyPDF2: {str(e)}")
                    # Se falhar com PyPDF2, tenta com pypdf
                    try:
                        file.seek(0)
                        pdf_reader = pypdf.PdfReader(file)
                        num_pages = len(pdf_reader.pages)
                        print(f"FileReadTool: PDF tem {num_pages} páginas (usando pypdf)")
                        
                        max_pages = min(20, num_pages)
                        print(f"FileReadTool: Processando {max_pages} páginas do PDF")
                        
                        for page_num in range(max_pages):
                            try:
                                page = pdf_reader.pages[page_num]
                                page_text = page.extract_text()
                                if page_text:
                                    # Limpa o texto da página
                                    page_text = clean_text(page_text)
                                    
                                    if len(page_text) > 50:
                                        text += f"\n\n=== Página {page_num + 1} ===\n\n{page_text}"
                                        print(f"FileReadTool: Texto extraído da página {page_num + 1}: {len(page_text)} caracteres")
                                    else:
                                        print(f"FileReadTool: Página {page_num + 1} ignorada por ter pouco conteúdo")
                                else:
                                    print(f"FileReadTool: Nenhum texto extraído da página {page_num + 1}")
                            except Exception as e:
                                print(f"FileReadTool: Erro ao extrair texto da página {page_num + 1}: {str(e)}")
                                continue
                    except Exception as e:
                        print(f"FileReadTool: Erro ao ler PDF com pypdf: {str(e)}")
                        raise
            
            if not text.strip():
                print("FileReadTool: Nenhum texto extraído do PDF")
                return "Error: No text extracted from PDF"
            
            # Limpa o texto final
            text = clean_text(text)
            
            print(f"FileReadTool: Texto extraído com sucesso do PDF. Tamanho: {len(text)} caracteres")
            return text
            
        except Exception as e:
            print(f"FileReadTool: Erro ao extrair texto do PDF: {str(e)}")
            return f"Error: Failed to extract text from PDF: {str(e)}"

    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            print(f"FileReadTool: Extraindo texto do DOCX: {file_path}")
            text = ""
            
            doc = docx.Document(file_path)
            
            # Extrai texto dos parágrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extrai texto das tabelas
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            if not text.strip():
                print("FileReadTool: Nenhum texto extraído do DOCX")
                return "Error: No text extracted from DOCX"
            
            print(f"FileReadTool: Texto extraído com sucesso do DOCX. Tamanho: {len(text)} caracteres")
            return text
            
        except Exception as e:
            print(f"FileReadTool: Erro ao extrair texto do DOCX: {str(e)}")
            return f"Error: Failed to extract text from DOCX: {str(e)}"

    def _extract_text_from_markdown(self, file_path: str) -> str:
        """Extract text from Markdown file."""
        try:
            print(f"FileReadTool: Extraindo texto do Markdown: {file_path}")
            
            # Lista de encodings para tentar (incluindo UTF-8 com BOM)
            encodings = ['utf-8-sig', 'utf-8', 'latin1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        text = file.read()
                    
                    if not text.strip():
                        print(f"FileReadTool: Nenhum texto extraído do Markdown com encoding {encoding}")
                        continue
                    
                    print(f"FileReadTool: Texto extraído com sucesso do Markdown usando encoding {encoding}. Tamanho: {len(text)} caracteres")
                    return text
                    
                except (UnicodeDecodeError, Exception) as e:
                    print(f"FileReadTool: Erro ao ler Markdown com encoding {encoding}: {str(e)}")
                    continue
            
            print("FileReadTool: Não foi possível ler o Markdown com nenhum encoding")
            return "Error: Failed to extract text from Markdown: Could not decode with any encoding"
            
        except Exception as e:
            print(f"FileReadTool: Erro ao extrair texto do Markdown: {str(e)}")
            return f"Error: Failed to extract text from Markdown: {str(e)}"

    def _extract_text_from_zip(self, file_path: str, max_chars: int, max_depth: int = 3, current_depth: int = 0) -> str:
        """Extract text from ZIP file with support for nested ZIPs."""
        try:
            print(f"FileReadTool: Extraindo texto do ZIP (nível {current_depth + 1}): {file_path}")
            text = ""
            temp_dir = None
            
            # Verifica profundidade máxima para evitar loops infinitos
            if current_depth >= max_depth:
                print(f"FileReadTool: Profundidade máxima atingida ({max_depth}). Parando recursão.")
                return f"Error: Maximum ZIP depth reached ({max_depth}). Stopping recursion."
            
            try:
                # Verifica se o arquivo ZIP existe
                if not os.path.exists(file_path):
                    print(f"FileReadTool: Arquivo ZIP não encontrado: {file_path}")
                    return f"Error: ZIP file not found: {file_path}"
                
                # Verifica se é um arquivo ZIP válido
                if not zipfile.is_zipfile(file_path):
                    print(f"FileReadTool: Arquivo não é um ZIP válido: {file_path}")
                    return f"Error: File is not a valid ZIP: {file_path}"
                
                # Cria diretório temporário
                temp_dir = tempfile.mkdtemp()
                print(f"FileReadTool: Diretório temporário criado: {temp_dir}")
                
                # Extrai arquivo ZIP
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    # Lista arquivos no ZIP
                    file_list = zip_ref.namelist()
                    print(f"FileReadTool: Arquivos no ZIP (nível {current_depth + 1}): {file_list}")
                    
                    zip_ref.extractall(temp_dir)
                print(f"FileReadTool: Arquivo ZIP extraído com sucesso (nível {current_depth + 1})")
                
                # Processa arquivos extraídos
                processed_files = 0
                nested_zips = []
                
                for root, _, files in os.walk(temp_dir):
                    for file in files:
                        file_path_extracted = os.path.join(root, file)
                        print(f"FileReadTool: Processando arquivo extraído: {file}")
                        
                        try:
                            # Verifica se é um arquivo compactado aninhado
                            if is_compressed_file(file_path_extracted):
                                compressed_type = get_compressed_file_type(file_path_extracted)
                                print(f"FileReadTool: Arquivo compactado aninhado encontrado ({compressed_type}): {file}")
                                
                                # Por enquanto, só processa ZIPs aninhados
                                if compressed_type == 'zip':
                                    nested_zips.append(file_path_extracted)
                                else:
                                    print(f"FileReadTool: Tipo de arquivo compactado não suportado: {compressed_type}")
                                    # Tenta extrair como arquivo normal (pode ser um arquivo de texto com extensão .zip)
                                    file_text = self._extract_text_from_file(file_path_extracted)
                                    if file_text and not file_text.startswith("Error:"):
                                        text += f"\n\n=== {file} ===\n\n{file_text}"
                                        print(f"FileReadTool: Texto extraído do arquivo {file}: {len(file_text)} caracteres")
                                        processed_files += 1
                                continue
                            
                            # Para outros tipos de arquivo, extrai texto normalmente
                            file_text = self._extract_text_from_file(file_path_extracted)
                            
                            if file_text and not file_text.startswith("Error:"):
                                text += f"\n\n=== {file} ===\n\n{file_text}"
                                print(f"FileReadTool: Texto extraído do arquivo {file}: {len(file_text)} caracteres")
                                processed_files += 1
                            else:
                                print(f"FileReadTool: Erro ao extrair texto do arquivo {file}: {file_text}")
                                
                        except Exception as e:
                            print(f"FileReadTool: Erro ao processar arquivo {file}: {str(e)}")
                            continue
                
                # Processa ZIPs aninhados recursivamente
                if nested_zips:
                    print(f"FileReadTool: Processando {len(nested_zips)} ZIP(s) aninhado(s)...")
                    for nested_zip_path in nested_zips:
                        try:
                            nested_text = self._extract_text_from_zip(
                                nested_zip_path, 
                                max_chars, 
                                max_depth, 
                                current_depth + 1
                            )
                            
                            if nested_text and not nested_text.startswith("Error:"):
                                nested_file_name = os.path.basename(nested_zip_path)
                                text += f"\n\n=== ZIP ANINHADO: {nested_file_name} ===\n\n{nested_text}"
                                print(f"FileReadTool: Texto extraído do ZIP aninhado {nested_file_name}: {len(nested_text)} caracteres")
                                processed_files += 1
                            else:
                                print(f"FileReadTool: Erro ao extrair texto do ZIP aninhado: {nested_text}")
                                
                        except Exception as e:
                            print(f"FileReadTool: Erro ao processar ZIP aninhado {nested_zip_path}: {str(e)}")
                            continue
                
                print(f"FileReadTool: Total de arquivos processados (nível {current_depth + 1}): {processed_files}")
                
                if not text.strip():
                    print(f"FileReadTool: Nenhum texto extraído do ZIP (nível {current_depth + 1})")
                    return "Error: No text extracted from ZIP"
                
                # Limita tamanho do texto
                if len(text) > max_chars:
                    text = text[:max_chars] + f"\n\n[Texto truncado em {max_chars} caracteres]"
                
                print(f"FileReadTool: Texto extraído com sucesso do ZIP (nível {current_depth + 1}). Tamanho: {len(text)} caracteres")
                
                return text
                
            finally:
                # Limpa diretório temporário
                if temp_dir and os.path.exists(temp_dir):
                    try:
                        shutil.rmtree(temp_dir)
                        print(f"FileReadTool: Diretório temporário removido: {temp_dir}")
                    except Exception as e:
                        print(f"FileReadTool: Erro ao remover diretório temporário: {str(e)}")
            
        except zipfile.BadZipFile as e:
            print(f"FileReadTool: Arquivo ZIP corrompido: {str(e)}")
            return f"Error: Corrupted ZIP file: {str(e)}"
        except Exception as e:
            print(f"FileReadTool: Erro ao extrair texto do ZIP: {str(e)}")
            return f"Error: Failed to extract text from ZIP: {str(e)}"

def is_compressed_file(file_path: str) -> bool:
    """Verifica se um arquivo é um arquivo compactado."""
    file_extension = os.path.splitext(file_path)[1].lower()
    compressed_extensions = ['.zip', '.rar', '.7z', '.tar.gz', '.tar.bz2', '.gz', '.bz2']
    return file_extension in compressed_extensions

def get_compressed_file_type(file_path: str) -> str:
    """Retorna o tipo de arquivo compactado."""
    file_extension = os.path.splitext(file_path)[1].lower()
    if file_extension == '.zip':
        return 'zip'
    elif file_extension == '.rar':
        return 'rar'
    elif file_extension == '.7z':
        return '7z'
    elif file_extension in ['.tar.gz', '.tgz']:
        return 'tar.gz'
    elif file_extension in ['.tar.bz2', '.tbz2']:
        return 'tar.bz2'
    elif file_extension == '.gz':
        return 'gz'
    elif file_extension == '.bz2':
        return 'bz2'
    else:
        return 'unknown' 